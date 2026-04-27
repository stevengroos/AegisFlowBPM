from fastapi import APIRouter, Depends, HTTPException, Request, UploadFile, File, Form
from sqlalchemy.orm import Session
from pydantic import BaseModel
import httpx # 🔥 Usaremos httpx para llamadas HTTP asíncronas a las APIs de IA
import json
import base64
import io
import PyPDF2
import docx
import pandas as pd
import re
from app.db.session import get_db
from app.models import models
from app.api import deps
from app.core.global_audit import log_global_event

router = APIRouter()

class AIGenerateRequest(BaseModel):
    form_id: int
    prompt: str

# El System Prompt es la "personalidad" y las "reglas" de la IA
SYSTEM_PROMPT = """
Eres un Arquitecto de Software Experto en AegisFlow (un BPM SaaS).
Tu tarea es convertir el requerimiento del usuario (prompt) en un JSON estricto que representa un formulario (Canvas).

REGLAS ESTRICTAS PARA EL JSON DE SALIDA:
1. Debes devolver ÚNICAMENTE un objeto JSON. Nada de explicaciones, ni bloques de markdown (```json).
2. El JSON debe tener esta estructura exacta:
{
  "sections": [
    { "title": "Nombre de la Sección", "columns": 2 }
  ],
  "fields": [
    {
      "section_title": "Nombre de la Sección a la que pertenece",
      "label": "Nombre visible del campo",
      "api_name": "nombre_sin_espacios_ni_mayusculas",
      "field_type": "text | textarea | number | date | select | checkbox | relation | url | file | image | subform",
      "required": true_o_false,
      "options": "Opcion1, Opcion2" (solo si field_type es 'select'),
      "subform_config": [ (solo si field_type es 'subform')
         { "label": "Columna 1", "type": "text | number | date | select" }
      ]
    }
  ]
}
3. Nunca inventes 'field_types' que no estén en la lista permitida.
4. Si el usuario pide un archivo, usa 'file'. Si pide foto, usa 'image'.
"""

# --- PROMPT PARA GENERAR FLUJOS DESDE IMÁGENES ---
BLUEPRINT_SYSTEM_PROMPT = """
Eres un Arquitecto de Procesos (BPM) experto.
Tu tarea es analizar la imagen de un diagrama de flujo de trabajo (dibujado a mano o en software) y convertirlo ESTRICTAMENTE en un objeto JSON.

Reglas de interpretación:
1. Identifica cada caja/círculo/rombo como un "Status" (Estado).
2. Identifica cada flecha que conecta las cajas como una "Transition" (Transición).
3. Intenta deducir la forma BPMN adecuada para cada estado: 'start' (inicio), 'end' (fin), 'gateway' (decisión/rombo), o 'task' (tarea/caja normal).

REGLAS ESTRICTAS PARA EL JSON DE SALIDA:
1. Devuelve ÚNICAMENTE un objeto JSON válido, sin markdown (```json).
2. Estructura exacta requerida:
# CÁMBIALO PARA QUE QUEDE ASÍ:
{
  "statuses": [
    { "id": "temp_1", "name": "Inicio", "is_initial": true, "bpmn_shape": "start" },
    { "id": "temp_2", "name": "Revisar", "is_initial": false, "bpmn_shape": "task" }
  ],
  "transitions": [
    { "name": "Enviar a revisión", "from_status_id": "temp_1", "to_status_id": "temp_2" } 
  ]
}
3. Usa IDs temporales (temp_1, temp_2) para que las transiciones sepan qué estados conectar.
"""

# --- NUEVO: PROMPT PARA GENERAR FLUJOS DESDE TEXTO ---
BLUEPRINT_TEXT_SYSTEM_PROMPT = """
Eres un Arquitecto de Procesos (BPM) experto.
Tu tarea es analizar la descripción en texto de un proceso o flujo de trabajo y convertirlo ESTRICTAMENTE en un objeto JSON.

Reglas de interpretación:
1. Identifica cada paso/etapa como un "Status" (Estado).
2. Identifica las acciones que mueven de una etapa a otra como una "Transition" (Transición).
3. Asigna la forma BPMN adecuada: 'start' (inicio), 'end' (fin), 'gateway' (decisión/rombo), o 'task' (caja normal).

REGLAS ESTRICTAS PARA EL JSON DE SALIDA:
1. Devuelve ÚNICAMENTE un objeto JSON válido, sin markdown (```json).
2. Estructura exacta requerida:
{
  "statuses": [
    { "id": "temp_1", "name": "Inicio", "is_initial": true, "bpmn_shape": "start" },
    { "id": "temp_2", "name": "Revisar", "is_initial": false, "bpmn_shape": "task" }
  ],
  "transitions": [
    { "name": "Enviar a revisión", "from_status_id": "temp_1", "to_status_id": "temp_2" } 
  ]
}
"""

def extract_text_from_file_bytes(file_bytes: bytes, mime_type: str) -> str:
    """Función ayudante para extraer texto de documentos sin repetir código."""
    extracted_text = ""
    if mime_type == 'application/pdf':
        pdf_reader = PyPDF2.PdfReader(io.BytesIO(file_bytes))
        for page in pdf_reader.pages:
            extracted_text += page.extract_text() + "\n"
    elif mime_type in ['application/msword', 'application/vnd.openxmlformats-officedocument.wordprocessingml.document']:
        doc = docx.Document(io.BytesIO(file_bytes))
        extracted_text = "\n".join([para.text for para in doc.paragraphs])
    elif mime_type in ['application/vnd.ms-excel', 'application/vnd.openxmlformats-officedocument.spreadsheetml.sheet']:
        df = pd.read_excel(io.BytesIO(file_bytes))
        extracted_text = df.to_string()
    else:
        raise HTTPException(status_code=400, detail="Formato de archivo no soportado.")
    return extracted_text

def clean_json_response(content_str: str) -> dict:
    """Utilidad para limpiar la respuesta en caso de que la IA (como Claude) devuelva markdown."""
    content_str = content_str.strip()
    if content_str.startswith("```json"):
        content_str = content_str.split("\n", 1)[1]
    if content_str.startswith("```"):
        content_str = content_str.split("\n", 1)[1]
    if content_str.endswith("```"):
        content_str = content_str.rsplit("\n", 1)[0]
    return json.loads(content_str.strip())

@router.post("/generate-form")
async def generate_form_with_ai(
    req: AIGenerateRequest,
    request: Request,
    db: Session = Depends(get_db),
    current_user: models.User = Depends(deps.get_current_user)
):
    # 1. VERIFICAR CONFIGURACIÓN DE IA EN LA EMPRESA (Multi-tenant)
    company = db.query(models.Company).filter(models.Company.id == current_user.company_id).first()
    
    active_provider = company.ai_active_provider
    api_key = company.ai_api_key

    if not active_provider or not api_key:
        raise HTTPException(
            status_code=400, 
            detail="La Inteligencia Artificial no está configurada. Por favor, ve a las Políticas Globales e ingresa tu API Key de OpenAI, Anthropic o Gemini."
        )

    # 2. ENRUTAR AL PROVEEDOR SELECCIONADO
    try:
        if active_provider == "openai":
            return await call_openai(api_key, req.prompt)
        elif active_provider == "anthropic":
            return await call_anthropic(api_key, req.prompt)
        elif active_provider == "gemini":
            return await call_gemini(api_key, req.prompt)
        else:
            raise ValueError("Proveedor de IA no soportado.")
            
    except Exception as e:
        print(f"Error AI: {e}")
        raise HTTPException(status_code=500, detail="Error de comunicación con la API de Inteligencia Artificial.")


# =================================================================
# CONECTORES DE GENERACIÓN DE FORMULARIOS
# =================================================================

async def call_openai(api_key: str, user_prompt: str):
    url = "https://api.openai.com/v1/chat/completions"
    headers = { "Authorization": f"Bearer {api_key}", "Content-Type": "application/json" }
    payload = {
        "model": "gpt-4o-mini", 
        "messages": [
            {"role": "system", "content": SYSTEM_PROMPT},
            {"role": "user", "content": f"Genera un formulario para: {user_prompt}"}
        ],
        "temperature": 0.3, 
        "response_format": { "type": "json_object" } 
    }
    async with httpx.AsyncClient(verify=False) as client:
        response = await client.post(url, headers=headers, json=payload, timeout=40.0)
        response.raise_for_status()
        return clean_json_response(response.json()["choices"][0]["message"]["content"])

async def call_anthropic(api_key: str, user_prompt: str):
    url = "https://api.anthropic.com/v1/messages"
    headers = {
        "x-api-key": api_key,
        "anthropic-version": "2023-06-01",
        "content-type": "application/json"
    }
    payload = {
        "model": "claude-3-haiku-20240307", # Modelo súper rápido de Anthropic
        "max_tokens": 4000,
        "system": SYSTEM_PROMPT,
        "messages": [
            {"role": "user", "content": f"Genera un formulario para: {user_prompt}. Devuelve SOLO el JSON."}
        ],
        "temperature": 0.3
    }
    async with httpx.AsyncClient(verify=False) as client:
        response = await client.post(url, headers=headers, json=payload, timeout=40.0)
        response.raise_for_status()
        return clean_json_response(response.json()["content"][0]["text"])

async def call_gemini(api_key: str, user_prompt: str):
    url = f"https://generativelanguage.googleapis.com/v1beta/models/gemini-1.5-flash:generateContent?key={api_key}"
    headers = { "Content-Type": "application/json" }
    payload = {
        "system_instruction": { "parts": [ {"text": SYSTEM_PROMPT} ] },
        "contents": [{ "parts": [{"text": f"Genera un formulario para: {user_prompt}"}] }],
        "generationConfig": {
            "responseMimeType": "application/json", # Gemini soporta forzar JSON nativamente
            "temperature": 0.3
        }
    }
    async with httpx.AsyncClient(verify=False) as client:
        response = await client.post(url, headers=headers, json=payload, timeout=40.0)
        response.raise_for_status()
        content_str = response.json()["candidates"][0]["content"]["parts"][0]["text"]
        return clean_json_response(content_str)


# =================================================================
# 🔥 FASE 3.3: PROCESAMIENTO INTELIGENTE DE DOCUMENTOS (IDP) 🔥
# =================================================================

@router.post("/extract-document")
async def extract_document_with_ai(
    request: Request,
    file: UploadFile = File(...),
    expected_fields: str = Form(...),
    db: Session = Depends(get_db),
    current_user: models.User = Depends(deps.get_current_user)
):
    company = db.query(models.Company).filter(models.Company.id == current_user.company_id).first()
    active_provider = company.ai_active_provider
    api_key = company.ai_api_key

    if not active_provider or not api_key:
        raise HTTPException(status_code=400, detail="La IA no está configurada. Ve a Políticas Globales para activar un proveedor.")

    file_bytes = await file.read()
    mime_type = file.content_type
    
    system_prompt = f"""
    Eres un asistente experto en extracción de datos (OCR inteligente).
    Tu tarea es extraer los datos del documento proporcionado y devolverlos EXACTAMENTE en este formato JSON, usando estas llaves:
    {expected_fields}
    
    Reglas:
    1. Devuelve SOLO un objeto JSON válido, sin bloques de código Markdown ni texto adicional.
    2. Si no encuentras un dato, devuelve un string vacío "".
    3. Si es un monto numérico, devuelve solo el número (ej. "1500.50").
    """

    try:
        # 1. SI ES UNA IMAGEN -> Usamos Visión
        if mime_type.startswith('image/'):
            base64_image = base64.b64encode(file_bytes).decode('utf-8')
            if active_provider == "openai":
                return await extract_with_openai_vision(api_key, system_prompt, base64_image, mime_type)
            elif active_provider == "anthropic":
                return await extract_with_anthropic_vision(api_key, system_prompt, base64_image, mime_type)
            elif active_provider == "gemini":
                return await extract_with_gemini_vision(api_key, system_prompt, base64_image, mime_type)
            
        # 2. SI ES UN DOCUMENTO -> Extraemos el texto primero y usamos IA de texto normal
        else:
            extracted_text = ""
            
            if mime_type == 'application/pdf':
                pdf_reader = PyPDF2.PdfReader(io.BytesIO(file_bytes))
                for page in pdf_reader.pages:
                    extracted_text += page.extract_text() + "\n"
            elif mime_type in ['application/msword', 'application/vnd.openxmlformats-officedocument.wordprocessingml.document']:
                doc = docx.Document(io.BytesIO(file_bytes))
                extracted_text = "\n".join([para.text for para in doc.paragraphs])
            elif mime_type in ['application/vnd.ms-excel', 'application/vnd.openxmlformats-officedocument.spreadsheetml.sheet']:
                df = pd.read_excel(io.BytesIO(file_bytes))
                extracted_text = df.to_string()
            else:
                raise HTTPException(status_code=400, detail="Formato de archivo no soportado para extracción.")

            if not extracted_text.strip():
                raise HTTPException(status_code=400, detail="No se pudo extraer texto del documento. Quizás es un PDF escaneado sin OCR.")

            # Enrutar el texto a la IA seleccionada
            if active_provider == "openai":
                return await extract_with_openai_text(api_key, system_prompt, extracted_text)
            elif active_provider == "anthropic":
                return await extract_with_anthropic_text(api_key, system_prompt, extracted_text)
            elif active_provider == "gemini":
                return await extract_with_gemini_text(api_key, system_prompt, extracted_text)

    except Exception as e:
        print(f"Error AI IDP: {e}")
        raise HTTPException(status_code=500, detail="Error al procesar el documento con la IA.")


# --- CONECTORES VISUALES Y DE TEXTO (IDP) ---

# OPENAI
async def extract_with_openai_vision(api_key: str, prompt: str, base64_img: str, mime_type: str):
    url = "https://api.openai.com/v1/chat/completions"
    headers = { "Authorization": f"Bearer {api_key}", "Content-Type": "application/json" }
    payload = {
        "model": "gpt-4o-mini", 
        "messages": [{"role": "user", "content": [{"type": "text", "text": prompt}, {"type": "image_url", "image_url": {"url": f"data:{mime_type};base64,{base64_img}"}}] }],
        "response_format": { "type": "json_object" }, "temperature": 0.1
    }
    async with httpx.AsyncClient(verify=False) as client:
        response = await client.post(url, headers=headers, json=payload, timeout=40.0)
        response.raise_for_status()
        return clean_json_response(response.json()["choices"][0]["message"]["content"])

async def extract_with_openai_text(api_key: str, prompt: str, document_text: str):
    url = "https://api.openai.com/v1/chat/completions"
    headers = { "Authorization": f"Bearer {api_key}", "Content-Type": "application/json" }
    payload = {
        "model": "gpt-4o-mini", 
        "messages": [{"role": "system", "content": prompt}, {"role": "user", "content": f"Documento:\n\n{document_text[:30000]}"}],
        "response_format": { "type": "json_object" }, "temperature": 0.1
    }
    async with httpx.AsyncClient(verify=False) as client:
        response = await client.post(url, headers=headers, json=payload, timeout=40.0)
        response.raise_for_status()
        return clean_json_response(response.json()["choices"][0]["message"]["content"])


# ANTHROPIC (CLAUDE)
async def extract_with_anthropic_vision(api_key: str, prompt: str, base64_img: str, mime_type: str):
    url = "https://api.anthropic.com/v1/messages"
    headers = { "x-api-key": api_key, "anthropic-version": "2023-06-01", "content-type": "application/json" }
    payload = {
        "model": "claude-3-haiku-20240307", "max_tokens": 4000, "system": prompt, "temperature": 0.1,
        "messages": [{
            "role": "user", 
            "content": [
                {"type": "image", "source": {"type": "base64", "media_type": mime_type, "data": base64_img}},
                {"type": "text", "text": "Extrae los datos. Devuelve SOLO un JSON."}
            ]
        }]
    }
    async with httpx.AsyncClient(verify=False) as client:
        response = await client.post(url, headers=headers, json=payload, timeout=40.0)
        response.raise_for_status()
        return clean_json_response(response.json()["content"][0]["text"])

async def extract_with_anthropic_text(api_key: str, prompt: str, document_text: str):
    url = "https://api.anthropic.com/v1/messages"
    headers = { "x-api-key": api_key, "anthropic-version": "2023-06-01", "content-type": "application/json" }
    payload = {
        "model": "claude-3-haiku-20240307", "max_tokens": 4000, "system": prompt, "temperature": 0.1,
        "messages": [{"role": "user", "content": f"Documento:\n\n{document_text[:30000]}\n\nExtrae los datos. Devuelve SOLO un JSON."}]
    }
    async with httpx.AsyncClient(verify=False) as client:
        response = await client.post(url, headers=headers, json=payload, timeout=40.0)
        response.raise_for_status()
        return clean_json_response(response.json()["content"][0]["text"])


# GOOGLE GEMINI
async def extract_with_gemini_vision(api_key: str, prompt: str, base64_img: str, mime_type: str):
    url = f"https://generativelanguage.googleapis.com/v1beta/models/gemini-1.5-flash:generateContent?key={api_key}"
    headers = { "Content-Type": "application/json" }
    payload = {
        "system_instruction": {"parts": [{"text": prompt}]},
        "contents": [{
            "parts": [
                {"text": "Extrae los datos de este documento."},
                {"inline_data": {"mime_type": mime_type, "data": base64_img}}
            ]
        }],
        "generationConfig": { "responseMimeType": "application/json", "temperature": 0.1 }
    }
    async with httpx.AsyncClient(verify=False) as client:
        response = await client.post(url, headers=headers, json=payload, timeout=40.0)
        response.raise_for_status()
        return clean_json_response(response.json()["candidates"][0]["content"]["parts"][0]["text"])

async def extract_with_gemini_text(api_key: str, prompt: str, document_text: str):
    url = f"https://generativelanguage.googleapis.com/v1beta/models/gemini-1.5-flash:generateContent?key={api_key}"
    headers = { "Content-Type": "application/json" }
    payload = {
        "system_instruction": {"parts": [{"text": prompt}]},
        "contents": [{"parts": [{"text": f"Documento:\n\n{document_text[:30000]}"}]}],
        "generationConfig": { "responseMimeType": "application/json", "temperature": 0.1 }
    }
    async with httpx.AsyncClient(verify=False) as client:
        response = await client.post(url, headers=headers, json=payload, timeout=40.0)
        response.raise_for_status()
        return clean_json_response(response.json()["candidates"][0]["content"]["parts"][0]["text"])
    
# =================================================================
# 🔥 GENERADOR DE FLUJOS (BLUEPRINTS) DESDE IMÁGENES 🔥
# =================================================================

# =================================================================
# 🔥 GENERADOR DE FLUJOS (BLUEPRINTS) DESDE ARCHIVOS (IMG/PDF) 🔥
# =================================================================
@router.post("/generate-blueprint/file")
async def generate_blueprint_from_file(
    request: Request,
    file: UploadFile = File(...),
    db: Session = Depends(get_db),
    current_user: models.User = Depends(deps.get_current_user)
):
    company = db.query(models.Company).filter(models.Company.id == current_user.company_id).first()
    active_provider = company.ai_active_provider
    api_key = company.ai_api_key

    if not active_provider or not api_key:
        raise HTTPException(status_code=400, detail="La IA no está configurada. Activa un proveedor en las Políticas Globales.")

    file_bytes = await file.read()
    mime_type = file.content_type

    try:
        # 1. Si es imagen (Diagrama dibujado a mano)
        if mime_type.startswith('image/'):
            base64_image = base64.b64encode(file_bytes).decode('utf-8')
            if active_provider == "openai": return await extract_with_openai_vision(api_key, BLUEPRINT_SYSTEM_PROMPT, base64_image, mime_type)
            elif active_provider == "anthropic": return await extract_with_anthropic_vision(api_key, BLUEPRINT_SYSTEM_PROMPT, base64_image, mime_type)
            elif active_provider == "gemini": return await extract_with_gemini_vision(api_key, BLUEPRINT_SYSTEM_PROMPT, base64_image, mime_type)
            
        # 2. Si es PDF/Doc (Manual de procedimientos)
        else:
            extracted_text = extract_text_from_file_bytes(file_bytes, mime_type)
            if not extracted_text.strip(): raise HTTPException(400, "No se pudo extraer texto del documento.")
            
            instruction = "Lee este manual de procedimientos y extrae un diagrama de flujo BPMN válido basándote en los pasos descritos:\n\n" + extracted_text
            
            if active_provider == "openai": return await extract_with_openai_text(api_key, BLUEPRINT_TEXT_SYSTEM_PROMPT, instruction)
            elif active_provider == "anthropic": return await extract_with_anthropic_text(api_key, BLUEPRINT_TEXT_SYSTEM_PROMPT, instruction)
            elif active_provider == "gemini": return await extract_with_gemini_text(api_key, BLUEPRINT_TEXT_SYSTEM_PROMPT, instruction)

    except Exception as e:
        print(f"Error AI Blueprint: {e}")
        raise HTTPException(status_code=500, detail="Error al procesar el diagrama/manual con la IA.")
    
# =================================================================
# 🔥 NUEVO: GENERADOR DE FORMULARIOS DESDE ARCHIVOS (PDF/IMG) 🔥
# =================================================================
@router.post("/generate-form/file")
async def generate_form_from_file(
    request: Request,
    file: UploadFile = File(...),
    db: Session = Depends(get_db),
    current_user: models.User = Depends(deps.get_current_user)
):
    company = db.query(models.Company).filter(models.Company.id == current_user.company_id).first()
    active_provider = company.ai_active_provider
    api_key = company.ai_api_key

    if not active_provider or not api_key:
        raise HTTPException(status_code=400, detail="La IA no está configurada.")

    file_bytes = await file.read()
    mime_type = file.content_type
    prompt_instruction = "Analiza este documento y diseña un formulario que permita capturar toda su información."

    try:
        if mime_type.startswith('image/'):
            base64_image = base64.b64encode(file_bytes).decode('utf-8')
            if active_provider == "openai": return await extract_with_openai_vision(api_key, SYSTEM_PROMPT + "\n" + prompt_instruction, base64_image, mime_type)
            elif active_provider == "anthropic": return await extract_with_anthropic_vision(api_key, SYSTEM_PROMPT + "\n" + prompt_instruction, base64_image, mime_type)
            elif active_provider == "gemini": return await extract_with_gemini_vision(api_key, SYSTEM_PROMPT + "\n" + prompt_instruction, base64_image, mime_type)
        else:
            extracted_text = extract_text_from_file_bytes(file_bytes, mime_type)
            if not extracted_text.strip(): raise HTTPException(400, "No se pudo extraer texto del documento.")
            
            if active_provider == "openai": return await extract_with_openai_text(api_key, SYSTEM_PROMPT, prompt_instruction + "\n\nDocumento:\n" + extracted_text)
            elif active_provider == "anthropic": return await extract_with_anthropic_text(api_key, SYSTEM_PROMPT, prompt_instruction + "\n\nDocumento:\n" + extracted_text)
            elif active_provider == "gemini": return await extract_with_gemini_text(api_key, SYSTEM_PROMPT, prompt_instruction + "\n\nDocumento:\n" + extracted_text)
    except Exception as e:
        raise HTTPException(status_code=500, detail=f"Error generando formulario: {str(e)}")


# =================================================================
# 🔥 NUEVO: GENERADOR DE BLUEPRINTS DESDE TEXTO 🔥
# =================================================================
class AIBlueprintTextRequest(BaseModel):
    prompt: str

@router.post("/generate-blueprint/text")
async def generate_blueprint_from_text(
    req: AIBlueprintTextRequest,
    db: Session = Depends(get_db),
    current_user: models.User = Depends(deps.get_current_user)
):
    company = db.query(models.Company).filter(models.Company.id == current_user.company_id).first()
    active_provider = company.ai_active_provider
    api_key = company.ai_api_key

    if not active_provider or not api_key:
        raise HTTPException(status_code=400, detail="La IA no está configurada.")

    try:
        instruction = f"Construye un flujo BPMN para el siguiente proceso:\n\n{req.prompt}"
        if active_provider == "openai": return await extract_with_openai_text(api_key, BLUEPRINT_TEXT_SYSTEM_PROMPT, instruction)
        elif active_provider == "anthropic": return await extract_with_anthropic_text(api_key, BLUEPRINT_TEXT_SYSTEM_PROMPT, instruction)
        elif active_provider == "gemini": return await extract_with_gemini_text(api_key, BLUEPRINT_TEXT_SYSTEM_PROMPT, instruction)
    except Exception as e:
        raise HTTPException(status_code=500, detail=f"Error generando blueprint: {str(e)}")